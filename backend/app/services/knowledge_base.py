from __future__ import annotations

import hashlib
import io
import json
import math
import re
from pathlib import Path
from typing import Any

from sqlalchemy import delete, func, select, text
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.knowledge import KnowledgeBase, KnowledgeChunk, KnowledgeDocument, KnowledgeEmbedding, KnowledgeRetrievalLog
from app.schemas.knowledge import (
    KnowledgeBaseCreateRequest,
    KnowledgeBaseListResponse,
    KnowledgeBaseRead,
    KnowledgeChunkRead,
    KnowledgeChunkUpdateRequest,
    KnowledgeDocumentCreateRequest,
    KnowledgeDocumentDetail,
    KnowledgeDocumentRead,
    KnowledgeDocumentUpdateRequest,
    KnowledgeReindexAllResponse,
    KnowledgeSearchItem,
    KnowledgeSearchRequest,
    KnowledgeSearchResponse,
)
from app.services.embedding_client import EmbeddingError, embed_texts
from app.services.knowledge_faiss import rebuild_faiss_index
from app.services.knowledge_faiss import search_faiss


def ensure_knowledge_fts(db: Session) -> None:
    _ensure_knowledge_schema(db)
    db.execute(
        text(
            """
            CREATE VIRTUAL TABLE IF NOT EXISTS knowledge_chunks_fts
            USING fts5(chunk_id UNINDEXED, document_id UNINDEXED, title, content, tags)
            """
        )
    )
    db.commit()
    _ensure_default_knowledge_base(db)


def _ensure_knowledge_schema(db: Session) -> None:
    KnowledgeBase.__table__.create(bind=db.get_bind(), checkfirst=True)
    KnowledgeDocument.__table__.create(bind=db.get_bind(), checkfirst=True)
    KnowledgeChunk.__table__.create(bind=db.get_bind(), checkfirst=True)
    KnowledgeEmbedding.__table__.create(bind=db.get_bind(), checkfirst=True)
    KnowledgeRetrievalLog.__table__.create(bind=db.get_bind(), checkfirst=True)
    _ensure_columns(
        db,
        "knowledge_documents",
        {
            "knowledge_base_id": "INTEGER DEFAULT 1",
            "chunking_strategy": "VARCHAR(40) DEFAULT 'paragraph'",
            "chunk_size": "INTEGER DEFAULT 900",
            "chunk_overlap": "INTEGER DEFAULT 120",
            "separators_json": "TEXT DEFAULT '[\"\\n\\n\", \"\\n\", \"。\", \"；\"]'",
        },
    )


def _ensure_columns(db: Session, table: str, columns: dict[str, str]) -> None:
    existing = {row[1] for row in db.execute(text(f"PRAGMA table_info({table})")).all()}
    for name, ddl in columns.items():
        if name not in existing:
            db.execute(text(f"ALTER TABLE {table} ADD COLUMN {name} {ddl}"))
    db.commit()


def _ensure_default_knowledge_base(db: Session) -> KnowledgeBase:
    item = db.get(KnowledgeBase, 1)
    if item:
        return item
    item = KnowledgeBase(id=1, name="默认知识库", description="默认投研资料库")
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


def list_knowledge_bases(db: Session) -> KnowledgeBaseListResponse:
    ensure_knowledge_fts(db)
    bases = db.scalars(select(KnowledgeBase).order_by(KnowledgeBase.id)).all()
    return KnowledgeBaseListResponse(items=[knowledge_base_read(db, item) for item in bases])


def create_knowledge_base(db: Session, payload: KnowledgeBaseCreateRequest) -> KnowledgeBaseRead:
    ensure_knowledge_fts(db)
    item = KnowledgeBase(
        name=payload.name.strip(),
        description=payload.description.strip(),
        chunking_strategy=payload.chunking_strategy,
        chunk_size=payload.chunk_size,
        chunk_overlap=payload.chunk_overlap,
        separators_json=json.dumps(_normalize_list(payload.separators), ensure_ascii=False),
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return knowledge_base_read(db, item)


def knowledge_base_read(db: Session, item: KnowledgeBase) -> KnowledgeBaseRead:
    count = db.scalar(select(func.count(KnowledgeDocument.id)).where(KnowledgeDocument.knowledge_base_id == item.id))
    return KnowledgeBaseRead(
        id=item.id,
        name=item.name,
        description=item.description,
        chunking_strategy=item.chunking_strategy,
        chunk_size=item.chunk_size,
        chunk_overlap=item.chunk_overlap,
        separators=_json_list(item.separators_json),
        document_count=int(count or 0),
        created_at=item.created_at,
        updated_at=item.updated_at,
    )


def create_document(db: Session, payload: KnowledgeDocumentCreateRequest) -> KnowledgeDocumentDetail:
    ensure_knowledge_fts(db)
    knowledge_base = db.get(KnowledgeBase, payload.knowledge_base_id) or _ensure_default_knowledge_base(db)
    document = KnowledgeDocument(
        knowledge_base_id=knowledge_base.id,
        title=payload.title.strip(),
        doc_type=payload.doc_type.strip() or "note",
        source=payload.source.strip() or "manual",
        content=payload.content,
        summary=_summarize(payload.content),
        symbols_json=json.dumps(_normalize_list(payload.symbols), ensure_ascii=False),
        tags_json=json.dumps(_normalize_list(payload.tags), ensure_ascii=False),
        metadata_json=json.dumps(payload.metadata, ensure_ascii=False),
        enabled=payload.enabled,
        published_at=payload.published_at,
        status="indexed",
        chunking_strategy=payload.chunking_strategy or knowledge_base.chunking_strategy,
        chunk_size=payload.chunk_size or knowledge_base.chunk_size,
        chunk_overlap=payload.chunk_overlap if payload.chunk_overlap is not None else knowledge_base.chunk_overlap,
        separators_json=json.dumps(
            _normalize_list(payload.separators) or _json_list(knowledge_base.separators_json),
            ensure_ascii=False,
        ),
    )
    db.add(document)
    db.flush()
    _replace_chunks(db, document)
    db.commit()
    db.refresh(document)
    rebuild_faiss_index(db)
    return document_detail(db, document)


def create_document_from_file(
    db: Session,
    filename: str,
    content_type: str,
    data: bytes,
    knowledge_base_id: int = 1,
    chunking_strategy: str = "",
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
    separators: list[str] | None = None,
) -> KnowledgeDocumentDetail:
    knowledge_base = db.get(KnowledgeBase, knowledge_base_id) or _ensure_default_knowledge_base(db)
    parsed = _parse_uploaded_file(filename, content_type, data)
    payload = KnowledgeDocumentCreateRequest(
        title=parsed["title"],
        content=parsed["content"],
        knowledge_base_id=knowledge_base.id,
        doc_type=parsed["doc_type"],
        source=f"upload:{filename}",
        symbols=[],
        tags=parsed["tags"],
        metadata=parsed["metadata"],
        chunking_strategy=chunking_strategy or knowledge_base.chunking_strategy,
        chunk_size=chunk_size or knowledge_base.chunk_size,
        chunk_overlap=chunk_overlap if chunk_overlap is not None else knowledge_base.chunk_overlap,
        separators=separators or _json_list(knowledge_base.separators_json),
    )
    return create_document(db, payload)


def list_documents(
    db: Session,
    q: str = "",
    limit: int = 50,
    knowledge_base_id: int | None = None,
) -> list[KnowledgeDocumentRead]:
    query = select(KnowledgeDocument).order_by(KnowledgeDocument.id.desc()).limit(min(max(limit, 1), 200))
    if knowledge_base_id:
        query = query.where(KnowledgeDocument.knowledge_base_id == knowledge_base_id)
    documents = db.scalars(query).all()
    if q.strip():
        needle = q.strip().lower()
        documents = [
            item
            for item in documents
            if needle in f"{item.title} {item.source} {item.doc_type} {item.tags_json} {item.symbols_json}".lower()
        ]
    return [document_read(item) for item in documents]


def get_document(db: Session, document_id: int) -> KnowledgeDocumentDetail | None:
    document = db.get(KnowledgeDocument, document_id)
    if not document:
        return None
    return document_detail(db, document)


def update_document(
    db: Session,
    document_id: int,
    payload: KnowledgeDocumentUpdateRequest,
) -> KnowledgeDocumentDetail | None:
    document = db.get(KnowledgeDocument, document_id)
    if not document:
        return None
    for field in ("title", "doc_type", "source", "published_at"):
        value = getattr(payload, field)
        if value is not None:
            setattr(document, field, value)
    if payload.symbols is not None:
        document.symbols_json = json.dumps(_normalize_list(payload.symbols), ensure_ascii=False)
    if payload.tags is not None:
        document.tags_json = json.dumps(_normalize_list(payload.tags), ensure_ascii=False)
    if payload.metadata is not None:
        document.metadata_json = json.dumps(payload.metadata, ensure_ascii=False)
    if payload.enabled is not None:
        document.enabled = payload.enabled
    db.add(document)
    db.commit()
    db.refresh(document)
    return document_detail(db, document)


def delete_document(db: Session, document_id: int) -> bool:
    ensure_knowledge_fts(db)
    document = db.get(KnowledgeDocument, document_id)
    if not document:
        return False
    db.execute(delete(KnowledgeEmbedding).where(KnowledgeEmbedding.document_id == document_id))
    db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.document_id == document_id))
    db.execute(text("DELETE FROM knowledge_chunks_fts WHERE document_id = :document_id"), {"document_id": document_id})
    db.delete(document)
    db.commit()
    rebuild_faiss_index(db)
    return True


def reindex_document(db: Session, document_id: int) -> KnowledgeDocumentDetail | None:
    ensure_knowledge_fts(db)
    document = db.get(KnowledgeDocument, document_id)
    if not document:
        return None
    _replace_chunks(db, document)
    document.status = "indexed"
    db.add(document)
    db.commit()
    db.refresh(document)
    rebuild_faiss_index(db)
    return document_detail(db, document)


def reindex_all_documents(db: Session) -> KnowledgeReindexAllResponse:
    ensure_knowledge_fts(db)
    documents = db.scalars(select(KnowledgeDocument).order_by(KnowledgeDocument.id)).all()
    reindexed = 0
    failed = 0
    for document in documents:
        try:
            _replace_chunks(db, document)
            document.status = "indexed"
            db.add(document)
            db.commit()
            db.refresh(document)
            reindexed += 1
        except Exception:
            db.rollback()
            failed += 1
            ensure_knowledge_fts(db)
    faiss = rebuild_faiss_index(db)
    return KnowledgeReindexAllResponse(
        total=len(documents),
        reindexed=reindexed,
        failed=failed,
        faiss=faiss,
    )


def search_knowledge(db: Session, payload: KnowledgeSearchRequest) -> KnowledgeSearchResponse:
    ensure_knowledge_fts(db)
    rows = _merge_search_rows(_fts_search(db, payload), _vector_search(db, payload))
    if not rows:
        rows = _fallback_search(db, payload)
    items = [_row_to_search_item(row) for row in rows[: payload.top_k]]
    db.add(
        KnowledgeRetrievalLog(
            query=payload.query,
            filters_json=json.dumps(
                {
                    "symbols": payload.symbols,
                    "doc_types": payload.doc_types,
                    "tags": payload.tags,
                    "top_k": payload.top_k,
                },
                ensure_ascii=False,
            ),
            result_json=json.dumps([item.model_dump(mode="json") for item in items], ensure_ascii=False),
        )
    )
    db.commit()
    return KnowledgeSearchResponse(query=payload.query, items=items)


def document_read(document: KnowledgeDocument) -> KnowledgeDocumentRead:
    return KnowledgeDocumentRead(
        id=document.id,
        knowledge_base_id=document.knowledge_base_id,
        title=document.title,
        doc_type=document.doc_type,
        source=document.source,
        summary=document.summary,
        symbols=_json_list(document.symbols_json),
        tags=_json_list(document.tags_json),
        metadata=_json_dict(document.metadata_json),
        status=document.status,
        enabled=document.enabled,
        chunk_count=document.chunk_count,
        chunking_strategy=document.chunking_strategy,
        chunk_size=document.chunk_size,
        chunk_overlap=document.chunk_overlap,
        separators=_json_list(document.separators_json),
        published_at=document.published_at,
        created_at=document.created_at,
        updated_at=document.updated_at,
    )


def update_chunk(
    db: Session,
    chunk_id: int,
    payload: KnowledgeChunkUpdateRequest,
) -> KnowledgeChunkRead | None:
    ensure_knowledge_fts(db)
    chunk = db.get(KnowledgeChunk, chunk_id)
    if not chunk:
        return None
    document = db.get(KnowledgeDocument, chunk.document_id)
    if not document:
        return None
    metadata = _json_dict(chunk.metadata_json)
    if payload.content is not None:
        chunk.content = payload.content
        chunk.summary = _summarize(payload.content, max_chars=120)
        chunk.token_count = max(len(payload.content) // 2, 1)
        chunk.content_hash = hashlib.sha256(payload.content.encode("utf-8")).hexdigest()
    if payload.tags is not None:
        tags = _normalize_list(payload.tags)
    else:
        tags = _json_list(metadata.get("tags_json", "[]"))
    metadata["tags"] = tags
    metadata["tags_json"] = json.dumps(tags, ensure_ascii=False)
    metadata["edited"] = True
    chunk.metadata_json = json.dumps(metadata, ensure_ascii=False)
    db.add(chunk)
    db.execute(text("DELETE FROM knowledge_chunks_fts WHERE chunk_id = :chunk_id"), {"chunk_id": chunk.id})
    db.execute(
        text(
            """
            INSERT INTO knowledge_chunks_fts(chunk_id, document_id, title, content, tags)
            VALUES (:chunk_id, :document_id, :title, :content, :tags)
            """
        ),
        {
            "chunk_id": chunk.id,
            "document_id": document.id,
            "title": document.title,
            "content": chunk.content,
            "tags": " ".join(tags),
        },
    )
    db.execute(delete(KnowledgeEmbedding).where(KnowledgeEmbedding.chunk_id == chunk.id))
    db.flush()
    _index_embeddings_for_chunks(db, [chunk])
    db.commit()
    db.refresh(chunk)
    rebuild_faiss_index(db)
    return chunk_read(chunk)


def document_detail(db: Session, document: KnowledgeDocument) -> KnowledgeDocumentDetail:
    chunks = db.scalars(
        select(KnowledgeChunk)
        .where(KnowledgeChunk.document_id == document.id)
        .order_by(KnowledgeChunk.chunk_index)
    ).all()
    base = document_read(document).model_dump()
    return KnowledgeDocumentDetail(
        **base,
        content=document.content,
        chunks=[chunk_read(chunk) for chunk in chunks],
    )


def chunk_read(chunk: KnowledgeChunk) -> KnowledgeChunkRead:
    metadata = _json_dict(chunk.metadata_json)
    return KnowledgeChunkRead(
        id=chunk.id,
        document_id=chunk.document_id,
        chunk_index=chunk.chunk_index,
        content=chunk.content,
        summary=chunk.summary,
        tags=_json_list(metadata.get("tags_json", "[]")),
        token_count=chunk.token_count,
        metadata=metadata,
    )


def _parse_uploaded_file(filename: str, content_type: str, data: bytes) -> dict[str, Any]:
    suffix = Path(filename).suffix.lower()
    title = Path(filename).stem.strip() or "未命名资料"
    if suffix in {".txt", ".md", ".csv", ".json", ".html", ".htm"} or content_type.startswith("text/"):
        content = _decode_text(data)
    elif suffix == ".pdf":
        content = _extract_pdf_text(data)
    elif suffix == ".docx":
        content = _extract_docx_text(data)
    else:
        content = _decode_text(data)
    content = content.strip()
    if not content:
        raise ValueError("Uploaded file has no extractable text.")
    return {
        "title": title,
        "content": content,
        "doc_type": _doc_type_from_suffix(suffix),
        "tags": _tags_from_filename(filename),
        "metadata": {
            "filename": filename,
            "content_type": content_type,
            "size_bytes": len(data),
            "parser": _parser_name(suffix, content_type),
        },
    }


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("PDF parsing requires pypdf. Install backend[rag].") from exc
    reader = PdfReader(io.BytesIO(data))
    return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("DOCX parsing requires python-docx. Install backend[rag].") from exc
    document = Document(io.BytesIO(data))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n\n".join(paragraphs)


def _doc_type_from_suffix(suffix: str) -> str:
    mapping = {
        ".pdf": "pdf",
        ".docx": "document",
        ".md": "markdown",
        ".txt": "text",
        ".csv": "table",
        ".json": "data",
    }
    return mapping.get(suffix, "document")


def _parser_name(suffix: str, content_type: str) -> str:
    if suffix == ".pdf":
        return "pypdf"
    if suffix == ".docx":
        return "python-docx"
    if suffix or content_type.startswith("text/"):
        return "text"
    return "text-fallback"


def _tags_from_filename(filename: str) -> list[str]:
    stem = Path(filename).stem
    tokens = re.findall(r"[\u4e00-\u9fff]{2,}|[A-Za-z0-9_]{2,}", stem)
    return tokens[:6]


def _chunk_tags(content: str, document: KnowledgeDocument) -> list[str]:
    tags: list[str] = []
    for item in _json_list(document.symbols_json) + _json_list(document.tags_json):
        if item not in tags:
            tags.append(item)
    for item in _keywords(content):
        if len(item) < 2 or item in tags:
            continue
        tags.append(item)
        if len(tags) >= 8:
            break
    return tags


def _replace_chunks(db: Session, document: KnowledgeDocument) -> None:
    db.execute(delete(KnowledgeEmbedding).where(KnowledgeEmbedding.document_id == document.id))
    db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.document_id == document.id))
    db.execute(text("DELETE FROM knowledge_chunks_fts WHERE document_id = :document_id"), {"document_id": document.id})
    chunks = _chunk_text(
        document.content,
        strategy=document.chunking_strategy,
        max_chars=document.chunk_size,
        overlap=document.chunk_overlap,
        separators=_json_list(document.separators_json),
    )
    for index, chunk_text in enumerate(chunks):
        chunk_tags = _chunk_tags(chunk_text, document)
        chunk = KnowledgeChunk(
            document_id=document.id,
            chunk_index=index,
            content=chunk_text,
            summary=_summarize(chunk_text, max_chars=120),
            token_count=max(len(chunk_text) // 2, 1),
            content_hash=hashlib.sha256(chunk_text.encode("utf-8")).hexdigest(),
            metadata_json=json.dumps(
                {
                    "chunking": document.chunking_strategy,
                    "tags": chunk_tags,
                    "tags_json": json.dumps(chunk_tags, ensure_ascii=False),
                },
                ensure_ascii=False,
            ),
        )
        db.add(chunk)
        db.flush()
        db.execute(
            text(
                """
                INSERT INTO knowledge_chunks_fts(chunk_id, document_id, title, content, tags)
                VALUES (:chunk_id, :document_id, :title, :content, :tags)
                """
            ),
            {
                "chunk_id": chunk.id,
                "document_id": document.id,
                "title": document.title,
                "content": chunk.content,
                "tags": " ".join(chunk_tags),
            },
        )
    document.chunk_count = len(chunks)
    document.summary = _summarize(document.content)
    db.add(document)
    db.flush()
    _index_embeddings(db, document.id)


def _index_embeddings(db: Session, document_id: int) -> None:
    chunks = db.scalars(
        select(KnowledgeChunk)
        .where(KnowledgeChunk.document_id == document_id)
        .order_by(KnowledgeChunk.chunk_index)
    ).all()
    if not chunks:
        return
    _index_embeddings_for_chunks(db, chunks)


def _index_embeddings_for_chunks(db: Session, chunks: list[KnowledgeChunk]) -> None:
    if not chunks:
        return
    try:
        vectors = embed_texts([chunk.content for chunk in chunks])
    except EmbeddingError:
        return
    for chunk, vector in zip(chunks, vectors):
        db.add(
            KnowledgeEmbedding(
                chunk_id=chunk.id,
                document_id=chunk.document_id,
                model=settings.embedding_model,
                dimension=len(vector),
                vector_json=json.dumps(vector),
            )
        )


def _fts_search(db: Session, payload: KnowledgeSearchRequest) -> list[dict[str, Any]]:
    terms = _fts_query(payload.query)
    if not terms:
        return []
    try:
        rows = db.execute(
            text(
                """
                SELECT c.id AS chunk_id, c.document_id, c.content, c.summary, c.metadata_json,
                       d.title, d.source, d.doc_type, d.symbols_json, d.tags_json,
                       bm25(knowledge_chunks_fts) AS rank
                FROM knowledge_chunks_fts
                JOIN knowledge_chunks c ON c.id = knowledge_chunks_fts.chunk_id
                JOIN knowledge_documents d ON d.id = c.document_id
                WHERE knowledge_chunks_fts MATCH :query AND d.enabled = 1
                ORDER BY rank
                LIMIT :limit
                """
            ),
            {"query": terms, "limit": max(payload.top_k * 4, 20)},
        ).mappings()
    except Exception:
        return []
    return _filter_rows([dict(row) for row in rows], payload)


def _fallback_search(db: Session, payload: KnowledgeSearchRequest) -> list[dict[str, Any]]:
    chunks = db.execute(
        text(
            """
            SELECT c.id AS chunk_id, c.document_id, c.content, c.summary, c.metadata_json,
                   d.title, d.source, d.doc_type, d.symbols_json, d.tags_json
            FROM knowledge_chunks c
            JOIN knowledge_documents d ON d.id = c.document_id
            WHERE d.enabled = 1
            LIMIT 500
            """
        )
    ).mappings()
    keywords = _keywords(payload.query)
    rows = []
    for row in chunks:
        data = dict(row)
        haystack = f"{data['title']} {data['content']} {' '.join(_json_list(data['tags_json']))}"
        score = sum(1 for keyword in keywords if keyword.lower() in haystack.lower())
        if score > 0:
            data["rank"] = -float(score)
            rows.append(data)
    return _filter_rows(sorted(rows, key=lambda item: item["rank"]), payload)


def _vector_search(db: Session, payload: KnowledgeSearchRequest) -> list[dict[str, Any]]:
    try:
        query_vector = embed_texts([payload.query])[0]
    except (EmbeddingError, IndexError):
        return []
    faiss_rows = _faiss_vector_search(db, payload, query_vector)
    if faiss_rows:
        return faiss_rows
    return _sqlite_vector_search(db, payload, query_vector)


def _faiss_vector_search(
    db: Session,
    payload: KnowledgeSearchRequest,
    query_vector: list[float],
) -> list[dict[str, Any]]:
    hits = search_faiss(query_vector, max(payload.top_k * 4, 20))
    if not hits:
        return []
    score_by_chunk_id = {int(hit["chunk_id"]): float(hit.get("vector_score") or 0.0) for hit in hits}
    rows = db.execute(
        select(
            KnowledgeChunk.id.label("chunk_id"),
            KnowledgeChunk.document_id,
            KnowledgeChunk.content,
            KnowledgeChunk.summary,
            KnowledgeChunk.metadata_json,
            KnowledgeDocument.title,
            KnowledgeDocument.source,
            KnowledgeDocument.doc_type,
            KnowledgeDocument.symbols_json,
            KnowledgeDocument.tags_json,
        )
        .join(KnowledgeDocument, KnowledgeDocument.id == KnowledgeChunk.document_id)
        .where(KnowledgeDocument.enabled.is_(True), KnowledgeChunk.id.in_(score_by_chunk_id))
    ).mappings()
    scored: list[dict[str, Any]] = []
    for row in rows:
        data = dict(row)
        similarity = score_by_chunk_id.get(int(data["chunk_id"]), 0.0)
        data["rank"] = -similarity
        data["vector_score"] = similarity
        scored.append(data)
    return _filter_rows(sorted(scored, key=lambda item: item["rank"]), payload)


def _sqlite_vector_search(
    db: Session,
    payload: KnowledgeSearchRequest,
    query_vector: list[float],
) -> list[dict[str, Any]]:
    rows = db.execute(
        text(
            """
            SELECT c.id AS chunk_id, c.document_id, c.content, c.summary, c.metadata_json,
                   d.title, d.source, d.doc_type, d.symbols_json, d.tags_json,
                   e.vector_json
            FROM knowledge_embeddings e
            JOIN knowledge_chunks c ON c.id = e.chunk_id
            JOIN knowledge_documents d ON d.id = c.document_id
            WHERE d.enabled = 1 AND e.model = :model
            LIMIT 1000
            """
        ),
        {"model": settings.embedding_model},
    ).mappings()
    scored: list[dict[str, Any]] = []
    for row in rows:
        data = dict(row)
        vector = _json_float_list(data.pop("vector_json", "[]"))
        if not vector:
            continue
        similarity = _cosine_similarity(query_vector, vector)
        data["rank"] = -similarity
        data["vector_score"] = similarity
        scored.append(data)
    return _filter_rows(sorted(scored, key=lambda item: item["rank"]), payload)


def _merge_search_rows(
    fts_rows: list[dict[str, Any]],
    vector_rows: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    merged: dict[int, dict[str, Any]] = {}
    for index, row in enumerate(fts_rows):
        item = dict(row)
        item["hybrid_score"] = item.get("hybrid_score", 0.0) + 1.0 / (index + 1)
        merged[int(item["chunk_id"])] = item
    for index, row in enumerate(vector_rows):
        chunk_id = int(row["chunk_id"])
        item = merged.get(chunk_id, dict(row))
        item["hybrid_score"] = item.get("hybrid_score", 0.0) + 1.2 / (index + 1)
        if "vector_score" in row:
            item["vector_score"] = row["vector_score"]
        merged[chunk_id] = item
    return sorted(merged.values(), key=lambda item: item.get("hybrid_score", 0.0), reverse=True)


def _filter_rows(rows: list[dict[str, Any]], payload: KnowledgeSearchRequest) -> list[dict[str, Any]]:
    result = []
    symbols = set(_normalize_list(payload.symbols))
    doc_types = set(_normalize_list(payload.doc_types))
    tags = set(_normalize_list(payload.tags))
    for row in rows:
        row_symbols = set(_json_list(row["symbols_json"]))
        row_tags = set(_json_list(row["tags_json"]))
        if symbols and not (symbols & row_symbols):
            continue
        if doc_types and row["doc_type"] not in doc_types:
            continue
        if tags and not (tags & row_tags):
            continue
        result.append(row)
    return result


def _row_to_search_item(row: dict[str, Any]) -> KnowledgeSearchItem:
    rank = float(row.get("rank", 0) or 0)
    if "hybrid_score" in row:
        score = round(min(float(row["hybrid_score"]), 1.0), 4)
    else:
        score = round(1 / (1 + max(rank, 0)), 4) if rank >= 0 else round(min(abs(rank), 10) / 10, 4)
    return KnowledgeSearchItem(
        chunk_id=int(row["chunk_id"]),
        document_id=int(row["document_id"]),
        title=str(row["title"]),
        snippet=_snippet(str(row["content"])),
        score=score,
        source=str(row["source"]),
        doc_type=str(row["doc_type"]),
        symbols=_json_list(row["symbols_json"]),
        tags=_json_list(row["tags_json"]),
        citation=f"doc:{row['document_id']}#chunk:{row['chunk_id']}",
        metadata=_json_dict(row.get("metadata_json", "{}")),
    )


def _chunk_text(
    content: str,
    strategy: str = "paragraph",
    max_chars: int = 900,
    overlap: int = 120,
    separators: list[str] | None = None,
) -> list[str]:
    if strategy == "characters":
        return _chunk_by_characters(content, max_chars=max_chars, overlap=overlap)
    if strategy == "separators":
        return _chunk_by_separators(content, separators or ["\n\n", "\n", "。", "；"], max_chars=max_chars, overlap=overlap)
    paragraphs = [item.strip() for item in re.split(r"\n\s*\n", content) if item.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs or [content.strip()]:
        if len(current) + len(paragraph) + 2 <= max_chars:
            current = f"{current}\n\n{paragraph}".strip()
            continue
        if current:
            chunks.append(current)
        prefix = current[-overlap:] if overlap and current else ""
        current = f"{prefix}\n\n{paragraph}".strip() if prefix else paragraph
    if current:
        chunks.append(current)
    return chunks or [content[:max_chars]]


def _chunk_by_characters(content: str, max_chars: int, overlap: int) -> list[str]:
    compact = content.strip()
    if not compact:
        return []
    max_chars = max(max_chars, 1)
    overlap = min(max(overlap, 0), max_chars - 1) if max_chars > 1 else 0
    chunks = []
    start = 0
    while start < len(compact):
        end = min(start + max_chars, len(compact))
        chunks.append(compact[start:end])
        if end >= len(compact):
            break
        start = max(end - overlap, start + 1)
    return chunks


def _chunk_by_separators(content: str, separators: list[str], max_chars: int, overlap: int) -> list[str]:
    escaped = [re.escape(item) for item in separators if item]
    if not escaped:
        return _chunk_by_characters(content, max_chars=max_chars, overlap=overlap)
    parts = [item.strip() for item in re.split("|".join(escaped), content) if item.strip()]
    chunks: list[str] = []
    current = ""
    for part in parts:
        if len(current) + len(part) + 1 <= max_chars:
            current = f"{current}\n{part}".strip()
            continue
        if current:
            chunks.append(current)
        prefix = current[-overlap:] if overlap and current else ""
        current = f"{prefix}\n{part}".strip() if prefix else part
    if current:
        chunks.append(current)
    return chunks or _chunk_by_characters(content, max_chars=max_chars, overlap=overlap)


def _summarize(content: str, max_chars: int = 180) -> str:
    compact = re.sub(r"\s+", " ", content).strip()
    return compact[:max_chars]


def _snippet(content: str, max_chars: int = 260) -> str:
    return _summarize(content, max_chars=max_chars)


def _fts_query(query: str) -> str:
    keywords = _keywords(query)
    return " OR ".join(keywords[:8])


def _keywords(query: str) -> list[str]:
    keywords: list[str] = []
    for token in re.findall(r"[\u4e00-\u9fff]+|[A-Za-z0-9_]+", query):
        if re.fullmatch(r"[\u4e00-\u9fff]+", token):
            if len(token) <= 4:
                keywords.append(token)
            keywords.extend(token[index : index + 2] for index in range(0, max(len(token) - 1, 0)))
        else:
            keywords.append(token)
    deduped: list[str] = []
    for keyword in keywords:
        if keyword and keyword not in deduped:
            deduped.append(keyword)
    return deduped


def _normalize_list(items: list[str]) -> list[str]:
    return [str(item).strip() for item in items if str(item).strip()]


def _json_list(value: str) -> list[str]:
    try:
        parsed = json.loads(value)
    except json.JSONDecodeError:
        return []
    return [str(item) for item in parsed] if isinstance(parsed, list) else []


def _json_dict(value: str) -> dict[str, Any]:
    try:
        parsed = json.loads(value)
    except json.JSONDecodeError:
        return {}
    return parsed if isinstance(parsed, dict) else {}


def _json_float_list(value: str) -> list[float]:
    try:
        parsed = json.loads(value)
    except json.JSONDecodeError:
        return []
    if not isinstance(parsed, list):
        return []
    return [float(item) for item in parsed]


def _cosine_similarity(left: list[float], right: list[float]) -> float:
    if len(left) != len(right) or not left:
        return 0.0
    dot = sum(a * b for a, b in zip(left, right))
    left_norm = math.sqrt(sum(a * a for a in left))
    right_norm = math.sqrt(sum(b * b for b in right))
    if left_norm == 0 or right_norm == 0:
        return 0.0
    return dot / (left_norm * right_norm)
